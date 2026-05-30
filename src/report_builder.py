"""
report_builder.py
-----------------
Builds a DOCX report matching the exact structure and style of the
Silver Oak University IEEE Student Branch report format (Alumni Connect style).

Structure:
  Cover Page  — A REPORT ON + Title + Date + Venue
  Introduction
  About the Speaker      (optional, bullet list)
  About the Session/Event (Date/Time/Venue/Participants + description)
  Conclusion
  SDG Impact             (optional, bullet list)
  IEEE Goals Achieved    (optional, bullet list)
  Acknowledgement        (optional)
"""

import re
from datetime import datetime
from pathlib  import Path

import fitz
from docx                   import Document
from docx.shared            import Pt, RGBColor, Inches, Cm
from docx.enum.text         import WD_ALIGN_PARAGRAPH
from docx.oxml.ns           import qn
from docx.oxml              import OxmlElement

from .config import DOCS_OUTPUT_DIR, OUTPUTS_DIR, PDF_OUTPUT_DIR


# ── Colours ────────────────────────────────────────────────────────────────────
BLACK      = RGBColor(0x00, 0x00, 0x00)
DARK_GREY  = RGBColor(0x22, 0x22, 0x22)
MID_GREY   = RGBColor(0x44, 0x44, 0x44)
LIGHT_GREY = RGBColor(0x88, 0x88, 0x88)
IEEE_BLUE  = RGBColor(0x00, 0x33, 0x99)

PDF_BLACK = (0, 0, 0)
PDF_GREY = (0.28, 0.28, 0.28)
PDF_LIGHT_GREY = (0.65, 0.65, 0.65)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_font(run, size_pt: float, bold=False, italic=False, color=BLACK):
    run.font.name        = "Calibri"
    run.font.size        = Pt(size_pt)
    run.bold             = bold
    run.italic           = italic
    run.font.color.rgb   = color


def _para_spacing(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)


def _add_section_heading(doc: Document, title: str):
    """Bold heading with colon, matching Alumni Connect style."""
    p = doc.add_paragraph()
    _para_spacing(p, before=14, after=4)
    r = p.add_run(title + ":")
    _set_font(r, 12, bold=True, color=BLACK)
    return p


def _add_body_para(doc: Document, text: str, indent=False):
    """Normal body paragraph."""
    if not text.strip():
        return
    p = doc.add_paragraph(text.strip())
    _para_spacing(p, before=0, after=5)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    for r in p.runs:
        _set_font(r, 11)
    return p


def _add_bullet(doc: Document, text: str):
    """Bullet point paragraph matching Alumni Connect style."""
    p = doc.add_paragraph(style="List Bullet")
    _para_spacing(p, before=1, after=3)
    r = p.add_run(text.strip())
    _set_font(r, 11)
    return p


def _add_labelled_line(doc: Document, label: str, value: str):
    """Bold label + normal value on one line — e.g. 'Date: 27th February 2026'."""
    p = doc.add_paragraph()
    _para_spacing(p, before=1, after=2)
    r_label = p.add_run(label + ": ")
    _set_font(r_label, 11, bold=True)
    r_value = p.add_run(value)
    _set_font(r_value, 11)
    return p


# ── Session metadata parser ────────────────────────────────────────────────────

_META_KEYS = ["date", "time", "venue", "participants", "location", "mode", "platform"]

def _parse_session_meta(text: str):
    """
    Split 'about the event' text into structured key-value pairs and
    remaining description paragraphs.
    Returns: (meta: list[(label,value)], body_lines: list[str])
    """
    meta       = []
    body_lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        m = re.match(r"^([A-Za-z ]{2,20})\s*:\s*(.+)$", line)
        if m and m.group(1).strip().lower() in _META_KEYS:
            meta.append((m.group(1).strip().title(), m.group(2).strip()))
        else:
            body_lines.append(line)
    return meta, body_lines


# ── Speaker bullet parser ──────────────────────────────────────────────────────

def _parse_speaker_bullets(text: str):
    """
    Parse speaker text into a list of bullet strings.
    If the text already has bullet markers (•, -, *), split on those.
    Otherwise treat each non-empty line as one bullet.
    """
    lines = [l.strip().lstrip("•-* ").strip() for l in text.splitlines() if l.strip()]
    return lines


# ── SDG / Goals bullet parser ─────────────────────────────────────────────────

def _parse_bullet_section(text: str):
    """Return (intro_sentence, bullets[])."""
    lines   = [l.strip() for l in text.splitlines() if l.strip()]
    intro   = ""
    bullets = []
    for line in lines:
        clean = line.lstrip("•-* ").strip()
        if line.startswith(("•", "-", "*")) or re.match(r"^SDG\s+\d+", line, re.I):
            bullets.append(clean)
        elif not intro:
            intro = clean
        else:
            # continuation of last bullet or new bullet
            if bullets:
                bullets[-1] += " " + clean
            else:
                bullets.append(clean)
    return intro, bullets


# ── Main builder ──────────────────────────────────────────────────────────────

def build_docx(
    event_name:  str,
    sections:    dict,
    output_path: str = None,
) -> Path:
    """
    Build a DOCX in the Alumni Connect report style.

    sections keys (all optional except event_name):
      introduction, about_the_speaker, about_the_event, description,
      conclusion, sdg_impact, ieee_goals, acknowledgement
    """
    doc = Document()

    # ── Page margins (matching the original) ──────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.18)
        sec.right_margin  = Inches(1.18)

    # ════════════════════════════════════════════════════════════════════════
    # COVER BLOCK
    # ════════════════════════════════════════════════════════════════════════

    # "A REPORT ON"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before=20, after=2)
    r = p.add_run("A REPORT ON")
    _set_font(r, 12, color=MID_GREY)

    # Event title — bold, large, centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p, before=2, after=16)
    r = p.add_run(event_name)
    _set_font(r, 18, bold=True, color=BLACK)

    # Date / Venue block (right-aligned, matching original layout)
    about_raw = sections.get("about_the_event", "")
    meta, _   = _parse_session_meta(about_raw)

    # Prefer explicit date/venue from meta; fallback to raw text
    date_val  = next((v for k, v in meta if k.lower() == "date"),  "")
    venue_val = next((v for k, v in meta if k.lower() == "venue"), "")

    if date_val or venue_val:
        for label, val in [("Date", date_val), ("Venue", venue_val)]:
            if val:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _para_spacing(p, before=1, after=2)
                r1 = p.add_run(label + ": ")
                _set_font(r1, 11, bold=True)
                r2 = p.add_run(val)
                _set_font(r2, 11)

    # Divider
    p = doc.add_paragraph()
    _para_spacing(p, before=18, after=4)
    r = p.add_run("─" * 80)
    _set_font(r, 8, color=LIGHT_GREY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ════════════════════════════════════════════════════════════════════════
    # INTRODUCTION
    # ════════════════════════════════════════════════════════════════════════
    intro_text = sections.get("introduction", "").strip()
    if intro_text:
        _add_section_heading(doc, "Introduction")
        for line in intro_text.splitlines():
            _add_body_para(doc, line)

    # ════════════════════════════════════════════════════════════════════════
    # ABOUT THE SPEAKER
    # ════════════════════════════════════════════════════════════════════════
    speaker_text = sections.get("about_the_speaker", "").strip()
    if speaker_text:
        _add_section_heading(doc, "About the Speaker")
        p = doc.add_paragraph()
        _para_spacing(p, before=0, after=3)
        r = p.add_run("The session was delivered by:")
        _set_font(r, 11)
        for bullet in _parse_speaker_bullets(speaker_text):
            _add_bullet(doc, bullet)

    # ════════════════════════════════════════════════════════════════════════
    # ABOUT THE SESSION / EVENT
    # ════════════════════════════════════════════════════════════════════════
    if about_raw.strip():
        _add_section_heading(doc, "About the Session")

        # Structured metadata lines (Date, Time, Venue, Participants…)
        for label, value in meta:
            _add_labelled_line(doc, label, value)

        if meta:
            doc.add_paragraph()   # small gap before description

        # Description paragraphs merged here (as in original)
        desc_text = sections.get("description", "").strip()
        combined  = "\n".join(_parse_session_meta(about_raw)[1])  # leftover body lines
        if desc_text:
            combined = (combined + "\n" + desc_text).strip()

        for line in combined.splitlines():
            _add_body_para(doc, line)

    # ════════════════════════════════════════════════════════════════════════
    # CONCLUSION
    # ════════════════════════════════════════════════════════════════════════
    conclusion_text = sections.get("conclusion", "").strip()
    if conclusion_text:
        _add_section_heading(doc, "Conclusion")
        for line in conclusion_text.splitlines():
            _add_body_para(doc, line)

    # ════════════════════════════════════════════════════════════════════════
    # SDG IMPACT  (optional)
    # ════════════════════════════════════════════════════════════════════════
    sdg_text = sections.get("sdg_impact", "").strip()
    if sdg_text:
        intro_s, bullets_s = _parse_bullet_section(sdg_text)
        p = doc.add_paragraph()
        _para_spacing(p, before=14, after=4)
        r1 = p.add_run("SDG Impact")
        _set_font(r1, 12, bold=True)
        if intro_s:
            r2 = p.add_run(": " + intro_s)
            _set_font(r2, 11)
        for b in bullets_s:
            _add_bullet(doc, b)

    # ════════════════════════════════════════════════════════════════════════
    # IEEE GOALS AND VISION ACHIEVED  (optional)
    # ════════════════════════════════════════════════════════════════════════
    goals_text = sections.get("ieee_goals", "").strip()
    if goals_text:
        intro_g, bullets_g = _parse_bullet_section(goals_text)
        _add_section_heading(doc, "IEEE Goals and Vision Achieved")
        if intro_g:
            _add_body_para(doc, intro_g)
        for b in bullets_g:
            _add_bullet(doc, b)

    # ════════════════════════════════════════════════════════════════════════
    # ACKNOWLEDGEMENT  (optional)
    # ════════════════════════════════════════════════════════════════════════
    ack_text = sections.get("acknowledgement", "").strip()
    if ack_text:
        for line in ack_text.splitlines():
            _add_body_para(doc, line)

    # ════════════════════════════════════════════════════════════════════════
    # SAVE
    # ════════════════════════════════════════════════════════════════════════
    DOCS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if output_path is None:
        safe  = "".join(c if c.isalnum() or c in " _-" else "_"
                        for c in event_name).strip().replace(" ", "_")
        ts    = datetime.now().strftime("%Y%m%d_%H%M")
        output_path = str(DOCS_OUTPUT_DIR / f"{safe}_{ts}.docx")

    doc.save(output_path)
    return Path(output_path)


def _safe_output_stem(event_name: str) -> str:
    safe = "".join(c if c.isalnum() or c in " _-" else "_" for c in event_name)
    return safe.strip().replace(" ", "_") or "event_report"


def _plain_section_lines(text: str) -> list[str]:
    return [line.strip().lstrip("•-* ").strip() for line in text.splitlines() if line.strip()]


def _pdf_add_textbox(page, rect, text, *, fontsize=11, bold=False, color=PDF_BLACK, align=0) -> float:
    fontname = "Helvetica-Bold" if bold else "Helvetica"
    spare = page.insert_textbox(
        rect,
        text,
        fontsize=fontsize,
        fontname=fontname,
        color=color,
        align=align,
    )
    used = rect.height - max(spare, 0)
    return max(used, fontsize + 4)


def build_pdf(
    event_name: str,
    sections: dict,
    output_path: str = None,
) -> Path:
    """
    Build a simple formal PDF report from the same generated report sections.
    """
    PDF_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if output_path is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        output_path = str(PDF_OUTPUT_DIR / f"{_safe_output_stem(event_name)}_{ts}.pdf")

    doc = fitz.open()
    width, height = 595, 842
    margin_x = 70
    margin_top = 68
    margin_bottom = 64
    line_gap = 8
    y = margin_top

    def new_page():
        return doc.new_page(width=width, height=height)

    page = new_page()

    def ensure_space(required: float):
        nonlocal page, y
        if y + required > height - margin_bottom:
            page = new_page()
            y = margin_top

    def add_text(text: str, *, fontsize=11, bold=False, color=PDF_BLACK, align=0, after=6):
        nonlocal y
        if not text.strip():
            return
        ensure_space(fontsize * 3)
        rect = fitz.Rect(margin_x, y, width - margin_x, height - margin_bottom)
        used = _pdf_add_textbox(page, rect, text.strip(), fontsize=fontsize, bold=bold, color=color, align=align)
        y += used + after

    add_text("A REPORT ON", fontsize=12, color=PDF_GREY, align=1, after=8)
    add_text(event_name, fontsize=18, bold=True, align=1, after=18)

    about_raw = sections.get("about_the_event", "")
    meta, _ = _parse_session_meta(about_raw)
    for label in ("Date", "Venue"):
        value = next((v for k, v in meta if k.lower() == label.lower()), "")
        if value:
            add_text(f"{label}: {value}", fontsize=11, bold=True, align=2, after=3)

    add_text("-" * 86, fontsize=8, color=PDF_LIGHT_GREY, align=1, after=12)

    section_order = [
        ("introduction", "Introduction"),
        ("about_the_speaker", "About the Speaker"),
        ("about_the_event", "About the Session"),
        ("conclusion", "Conclusion"),
        ("sdg_impact", "SDG Impact"),
        ("ieee_goals", "IEEE Goals and Vision Achieved"),
        ("acknowledgement", "Acknowledgement"),
    ]

    for key, heading in section_order:
        text = sections.get(key, "").strip()
        if not text:
            continue

        ensure_space(70)
        add_text(f"{heading}:", fontsize=12, bold=True, after=4)

        if key == "about_the_event":
            meta, body_lines = _parse_session_meta(text)
            for label, value in meta:
                add_text(f"{label}: {value}", fontsize=11, bold=True, after=2)
            for line in body_lines:
                add_text(line, fontsize=11, after=5)
            continue

        if key in {"about_the_speaker", "sdg_impact", "ieee_goals"}:
            for line in _plain_section_lines(text):
                add_text(f"- {line}", fontsize=11, after=3)
            y += line_gap
            continue

        for line in _plain_section_lines(text):
            add_text(line, fontsize=11, after=5)

    doc.save(output_path)
    doc.close()
    return Path(output_path)
